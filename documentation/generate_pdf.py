"""
Genere Panda_Rapport_de_Stage.pdf depuis le .docx via ReportLab
Tables et paragraphes sont rendus dans l'ordre exact du document.
"""
from docx import Document
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, KeepTogether,
)
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
import os, re

DOCX = r"C:\Users\Pro\Desktop\PFE O1\documentation\Panda_Rapport_de_Stage.docx"
PDF  = r"C:\Users\Pro\Desktop\PFE O1\documentation\Panda_Rapport_de_Stage.pdf"

# ── Styles ────────────────────────────────────────────────────────────────
BASE = getSampleStyleSheet()

def S(name, **kw):
    return ParagraphStyle(name, parent=BASE['Normal'], **kw)

sNormal  = S('sNormal',  fontSize=11, leading=15, alignment=TA_JUSTIFY,
             spaceAfter=4, fontName='Helvetica')
sH1      = S('sH1',      fontSize=16, leading=20, textColor=colors.HexColor('#003466'),
             fontName='Helvetica-Bold', spaceBefore=16, spaceAfter=6)
sH2      = S('sH2',      fontSize=13, leading=17, textColor=colors.HexColor('#00539C'),
             fontName='Helvetica-Bold', spaceBefore=12, spaceAfter=4)
sH3      = S('sH3',      fontSize=11, leading=14, textColor=colors.HexColor('#0072C6'),
             fontName='Helvetica-Bold', spaceBefore=8,  spaceAfter=3)
sBullet  = S('sBullet',  fontSize=10.5, leading=14, leftIndent=16,
             fontName='Helvetica', spaceAfter=2)
sCode    = S('sCode',    fontSize=8.5, leading=11.5, fontName='Courier',
             leftIndent=10, rightIndent=10,
             backColor=colors.HexColor('#F0F4F8'),
             textColor=colors.HexColor('#004600'),
             spaceBefore=3, spaceAfter=3)
sNote    = S('sNote',    fontSize=9.5, leading=13, fontName='Helvetica-Oblique',
             textColor=colors.HexColor('#555555'), leftIndent=12, spaceAfter=3)
sTitle   = S('sTitle',   fontSize=30, leading=36, fontName='Helvetica-Bold',
             textColor=colors.HexColor('#0078D4'), alignment=TA_CENTER, spaceAfter=10)
sSubtitle= S('sSubtitle',fontSize=14, leading=18, fontName='Helvetica',
             textColor=colors.HexColor('#333333'), alignment=TA_CENTER, spaceAfter=8)
sCenterB = S('sCenterB', fontSize=11, leading=15, fontName='Helvetica-Bold',
             textColor=colors.HexColor('#003466'), alignment=TA_CENTER, spaceAfter=3)
sCenter  = S('sCenter',  fontSize=11, leading=15, fontName='Helvetica',
             alignment=TA_CENTER, spaceAfter=3)

def esc(text):
    return (str(text)
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;'))

def is_code_para(p):
    """Detect code block: Courier New font or shading fill F0F4F8."""
    for run in p.runs:
        fn = run.font.name or ''
        if 'Courier' in fn:
            return True
    # Check paragraph shading fill via XML
    pPr = p._p.find(qn('w:pPr'))
    if pPr is not None:
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'), '')
            if fill.upper() in ('F0F4F8', 'EBF5FB'):
                return True
    return False

def para_to_flowable(p):
    text = p.text.strip()
    style_name = p.style.name if p.style else ''

    # Empty paragraph → small spacer
    if not text:
        return Spacer(1, 4)

    safe = esc(text)

    # Bold inline (first run bold check for mixed paragraphs)
    bold_runs = []
    for run in p.runs:
        if run.bold and run.text.strip():
            bold_runs.append(esc(run.text))

    if 'Heading 1' in style_name:
        return [
            HRFlowable(width="100%", thickness=1.2,
                       color=colors.HexColor('#003466'), spaceAfter=4),
            Paragraph(safe, sH1),
        ]
    if 'Heading 2' in style_name:
        return Paragraph(safe, sH2)
    if 'Heading 3' in style_name:
        return Paragraph(safe, sH3)
    if 'List Bullet' in style_name or 'List Number' in style_name:
        return Paragraph(f'&#x2022;&#160;&#160;{safe}', sBullet)
    if is_code_para(p):
        return Paragraph(safe, sCode)
    if 'Title' in style_name:
        return Paragraph(safe, sTitle)
    if style_name in ('Subtitle',):
        return Paragraph(safe, sSubtitle)

    # Center-aligned paragraphs (page de garde)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        # Check if bold dominant
        all_bold = all(r.bold for r in p.runs if r.text.strip())
        return Paragraph(safe, sCenterB if all_bold else sCenter)

    return Paragraph(safe, sNormal)

def table_to_flowable(tbl):
    rows_data = []
    for row in tbl.rows:
        row_cells = []
        for cell in row.cells:
            cell_text = ' '.join(cell.text.split())[:120]
            row_cells.append(cell_text)
        rows_data.append(row_cells)

    if not rows_data:
        return Spacer(1, 4)

    n_cols = max(len(r) for r in rows_data)
    # Deduplicate merged cells (docx repeats merged cell text)
    cleaned = []
    for row in rows_data:
        seen = []
        prev = None
        for cell in row:
            if cell != prev:
                seen.append(cell)
                prev = cell
            else:
                seen.append('')
        # pad/trim to n_cols
        while len(seen) < n_cols:
            seen.append('')
        cleaned.append(seen[:n_cols])

    avail_w = A4[0] - 5 * cm
    col_w   = avail_w / n_cols

    t = Table(cleaned, colWidths=[col_w] * n_cols, repeatRows=1,
              splitByRow=True)
    t.setStyle(TableStyle([
        # Header row
        ('BACKGROUND',    (0, 0), (-1, 0), colors.HexColor('#003466')),
        ('TEXTCOLOR',     (0, 0), (-1, 0), colors.white),
        ('FONTNAME',      (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE',      (0, 0), (-1, 0), 8.5),
        ('ALIGN',         (0, 0), (-1, 0), 'CENTER'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 5),
        # Body rows
        ('FONTNAME',      (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE',      (0, 1), (-1, -1), 8),
        ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
        ('GRID',          (0, 0), (-1, -1), 0.4, colors.HexColor('#AAAAAA')),
        ('TOPPADDING',    (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 3),
        ('LEFTPADDING',   (0, 0), (-1, -1), 4),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 4),
        ('ROWBACKGROUNDS',(0, 1), (-1, -1),
         [colors.HexColor('#EBF5FB'), colors.white]),
        ('WORDWRAP',      (0, 0), (-1, -1), True),
    ]))
    return t

# ── Build story in document order ────────────────────────────────────────
doc   = Document(DOCX)
story = []

# Iterate the body XML to preserve exact paragraph+table order
from docx.oxml.ns import qn
body = doc.element.body

para_idx  = 0
table_idx = 0
paras     = doc.paragraphs
tables    = doc.tables

for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

    if tag == 'p':
        if para_idx < len(paras):
            p = paras[para_idx]
            para_idx += 1
            fl = para_to_flowable(p)
            if isinstance(fl, list):
                story.extend(fl)
            else:
                story.append(fl)

    elif tag == 'tbl':
        if table_idx < len(tables):
            t = tables[table_idx]
            table_idx += 1
            story.append(Spacer(1, 4))
            story.append(table_to_flowable(t))
            story.append(Spacer(1, 6))

    elif tag == 'sectPr':
        pass  # section properties — ignore

# ── Page numbering footer ─────────────────────────────────────────────────
def footer_canvas(canvas, doc):
    canvas.saveState()
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(colors.HexColor('#666666'))
    page_num = f"Panda — Rapport de PFE — Ayoub Elmernissi — Page {doc.page}"
    canvas.drawCentredString(A4[0] / 2, 1.5 * cm, page_num)
    canvas.setStrokeColor(colors.HexColor('#003466'))
    canvas.setLineWidth(0.5)
    canvas.line(2.5 * cm, 1.8 * cm, A4[0] - 2.5 * cm, 1.8 * cm)
    canvas.restoreState()

# ── Render ────────────────────────────────────────────────────────────────
pdf_doc = SimpleDocTemplate(
    PDF, pagesize=A4,
    leftMargin=2.5*cm, rightMargin=2.5*cm,
    topMargin=2.5*cm,  bottomMargin=2.2*cm,
    title="Panda — Rapport de PFE",
    author="Ayoub Elmernissi",
    subject="Marketplace Freelance Full-Stack Laravel 12 / React 19",
    creator="ReportLab / Panda PDF Generator",
)

pdf_doc.build(story, onFirstPage=footer_canvas, onLaterPages=footer_canvas)

size_kb = os.path.getsize(PDF) // 1024
print(f"OK - PDF cree : {PDF} ({size_kb} KB)")
