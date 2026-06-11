# -*- coding: utf-8 -*-
"""
render_docx.py — render the shared block list to a classic-BTS-styled Word document.
Times New Roman body, Times New Roman blue headings (Heading 1/2/3 → native TOC), Word TOC +
Table of figures fields (auto page numbers on open), captions with SEQ fields,
page numbers in footer.
"""
import os
from PIL import Image as PILImage
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import theme as T
from inlines import parse_runs

CONTENT_W_CM = 16.6
MAX_IMG_H_CM = 20.5


def _rgb(hexs):
    h = hexs.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ── low-level XML helpers ─────────────────────────────────────────────────────
def add_field(paragraph, instr, cached=""):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = cached
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def add_seq(paragraph, name):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" SEQ {name} \\* ARABIC ")
    r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "1"; r.append(t); fld.append(r)
    paragraph._p.append(fld)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor.lstrip("#"))
    tcPr.append(shd)


def set_table_borders(table, color, sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color.lstrip("#"))
        borders.append(e)
    tblPr.append(borders)


def cell_left_bar(cell, color, sz=18):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    e = OxmlElement("w:left")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz)); e.set(qn("w:space"), "0")
    e.set(qn("w:color"), color.lstrip("#"))
    borders.append(e)
    tcPr.append(borders)


def enable_update_fields(doc):
    settings = doc.settings.element
    el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true")
    settings.append(el)


def add_page_border(section):
    sectPr = section._sectPr
    pg = OxmlElement("w:pgBorders")
    pg.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "18"); e.set(qn("w:space"), "24")
        e.set(qn("w:color"), T.TITLE_BLUE.lstrip("#"))
        pg.append(e)
    sectPr.append(pg)


def add_runs(paragraph, text, base_size=11.5, color=None, base_font="Times New Roman"):
    for t, b, i, m in parse_runs(text):
        run = paragraph.add_run(t)
        run.font.size = Pt(base_size)
        if m:
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1.5)
            run.font.color.rgb = _rgb("#9C2D41")
        else:
            run.font.name = base_font
            if color:
                run.font.color.rgb = _rgb(color)
        run.bold = b
        run.italic = i


# ── styles ────────────────────────────────────────────────────────────────────
def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = _rgb(T.BODY_BLACK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    def head(name, font, size, color, bold=True, align=None, before=12, after=6, indent=0):
        st = doc.styles[name]
        st.font.name = font; st.font.size = Pt(size); st.font.bold = bold
        st.font.color.rgb = _rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        if align is not None:
            st.paragraph_format.alignment = align
        if indent:
            st.paragraph_format.left_indent = Cm(indent)
        st.font.italic = False

    head("Heading 1", "Times New Roman", 18, T.TITLE_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=14)
    head("Heading 2", "Times New Roman", 14, T.HEADING_BLUE, before=14, after=6)
    head("Heading 3", "Times New Roman", 12.5, T.SUBHEAD_BLUE, before=10, after=4, indent=0.5)
    head("Heading 4", "Times New Roman", 11.5, "#333333", before=8, after=3, indent=0.9)
    doc.styles["Heading 4"].font.italic = True

    # caption
    if "Caption" in [s.name for s in doc.styles]:
        cap = doc.styles["Caption"]
    else:
        from docx.enum.style import WD_STYLE_TYPE
        cap = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = "Times New Roman"; cap.font.size = Pt(9.5); cap.font.italic = True
    cap.font.color.rgb = _rgb(T.CAPTION_GRAY); cap.font.bold = False
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3); cap.paragraph_format.space_after = Pt(12)


def fm_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(20); r.bold = True
    r.font.color.rgb = _rgb(T.TITLE_BLUE)
    return p


def footer_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.text = ""
    add_field(p, " PAGE ", "1")
    for r in p.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(11)


# ── block renderers ───────────────────────────────────────────────────────────
def add_image(doc, path):
    iw, ih = PILImage.open(path).size
    aspect = iw / ih
    w = CONTENT_W_CM
    h = w / aspect
    if h > MAX_IMG_H_CM:
        h = MAX_IMG_H_CM
        w = h * aspect
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(w))


def add_caption(doc, label, n, text):
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{label} ")
    add_seq(p, label)
    p.add_run(f": {text}")


def add_table(doc, headers, rows, n, cap, widths=None):
    ncol = len(headers) if headers else len(rows[0])
    table = doc.add_table(rows=0, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, T.TABLE_LINE, sz=4)
    if headers:
        hc = table.add_row().cells
        for j, htext in enumerate(headers):
            shade(hc[j], T.TABLE_HEAD)
            para = hc[j].paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = para.add_run(htext); r.font.name = "Times New Roman"; r.font.size = Pt(9.5)
            r.bold = True; r.font.color.rgb = _rgb("#FFFFFF")
    for ri, row in enumerate(rows):
        rc = table.add_row().cells
        for j, val in enumerate(row):
            if ri % 2 == 1:
                shade(rc[j], T.TABLE_ALT)
            para = rc[j].paragraphs[0]
            para.paragraph_format.space_after = Pt(1)
            add_runs(para, str(val), base_size=9.5)
    if widths:
        for j, wfrac in enumerate(widths):
            for cell in table.columns[j].cells:
                cell.width = Cm(CONTENT_W_CM * wfrac)
    if cap:
        add_caption(doc, "Tableau", n, cap)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc, cap, txt):
    if cap:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(cap); r.font.name = "Times New Roman"; r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = _rgb(T.HEADING_BLUE)
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, "#F5F7FA")
    set_table_borders(table, "#C5CDD8", sz=4)
    cell.paragraphs[0].text = ""
    # add generous left padding via cell margin
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("left", "200"), ("right", "120"), ("top", "100"), ("bottom", "100")]:
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), val); e.set(qn("w:type"), "dxa")
        tcMar.append(e)
    tcPr.append(tcMar)
    first = True
    for line in txt.split("\n"):
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        r = para.add_run(line if line else " ")
        r.font.name = "Consolas"; r.font.size = Pt(8.5); r.font.color.rgb = _rgb("#1B2430")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, "#EAF1F9")
    cell_left_bar(cell, T.HEADING_BLUE, sz=24)
    para = cell.paragraphs[0]
    add_runs(para, text, base_size=11)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc, items, level=0):
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8 + 0.7 * level)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        p.paragraph_format.space_after = Pt(2)
        mark = "•  " if level == 0 else "–  "
        mr = p.add_run(mark); mr.font.color.rgb = _rgb(T.HEADING_BLUE); mr.bold = True
        mr.font.name = "Times New Roman"
        add_runs(p, it)


def _cpara(doc, text, size=12, bold=False, italic=False, color=None, sa=4, sb=0, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Times New Roman"; r.font.size = Pt(size)
        r.bold = bold; r.italic = italic
        if color:
            r.font.color.rgb = _rgb(color)
    return p


def cover(doc):
    cov = T.COVER

    # Royal header
    _cpara(doc, "Royaume du Maroc", size=11, sa=2)
    _cpara(doc, cov["ministere"], size=10, italic=True, sa=10)

    # Institution bordered box
    tb = doc.add_table(rows=1, cols=1)
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tb, T.TITLE_BLUE, sz=8)
    cell = tb.cell(0, 0)
    cell.width = Cm(13)
    first = True
    for line in [cov["academie"], cov["direction"], cov["centre"]]:
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(5)
        para.paragraph_format.space_before = Pt(5)
        rr = para.add_run(line)
        rr.font.name = "Times New Roman"; rr.font.size = Pt(12)

    # Spacer
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(36)

    # RAPPORT DE PFE
    _cpara(doc, cov["report_kind"], size=16, bold=True, color=T.HEADING_BLUE, sa=20)

    # Project title (bold, dark)
    _cpara(doc, cov["project_title"], size=17, bold=True, sa=32)

    # Réalisé par / Encadrant table (no border)
    tb2 = doc.add_table(rows=1, cols=2)
    tb2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (lbl, vals) in enumerate([
            ("Réalisé par :", cov["students"]),
            ("Encadrant :", [cov["encadrant"]])]):
        c = tb2.cell(0, idx)
        c.width = Cm(8)
        lp = c.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_after = Pt(4)
        lr = lp.add_run(lbl)
        lr.font.name = "Times New Roman"; lr.font.size = Pt(12); lr.bold = True
        for v in vals:
            vp = c.add_paragraph()
            vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            vp.paragraph_format.space_after = Pt(2)
            vr = vp.add_run(v)
            vr.font.name = "Times New Roman"; vr.font.size = Pt(12)

    sp2 = doc.add_paragraph(); sp2.paragraph_format.space_after = Pt(30)

    # Centre de formation
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Centre de formation :"); r.font.name = "Times New Roman"; r.font.size = Pt(12); r.bold = True
    _cpara(doc, cov["lycee"], size=12, sa=40)

    # Year
    _cpara(doc, cov["year"], size=13, bold=True, sa=0)
    doc.add_page_break()


# ── main render ───────────────────────────────────────────────────────────────
def render(blocks, out_path):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7); section.page_width = Cm(21)
    section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(1.8)
    setup_styles(doc)
    footer_page_number(section)

    prev_break = False

    for b in blocks:
        k = b[0]
        if k == "cover":
            cover(doc)
        elif k == "frontmatter":
            fm_title(doc, b[1])
            for para in b[2]:
                pp = doc.add_paragraph(); add_runs(pp, para)
            doc.add_page_break()
        elif k == "toc":
            fm_title(doc, "Sommaire")
            p = doc.add_paragraph()
            add_field(p, ' TOC \\o "1-3" \\h \\z \\u ', "Le sommaire se met à jour à l'ouverture (clic droit > Mettre à jour les champs).")
            doc.add_page_break()
        elif k == "lof":
            fm_title(doc, "Table des figures")
            p = doc.add_paragraph()
            add_field(p, ' TOC \\h \\z \\c "Figure" ', "La table des figures se met à jour à l'ouverture.")
            doc.add_page_break()
        elif k == "lot":
            fm_title(doc, "Liste des tableaux")
            p = doc.add_paragraph()
            add_field(p, ' TOC \\h \\z \\c "Tableau" ', "La liste des tableaux se met à jour à l'ouverture.")
            doc.add_page_break()
        elif k == "h1":
            if not prev_break:
                doc.add_page_break()
            doc.add_paragraph(b[1], style="Heading 1")
        elif k == "h2":
            doc.add_paragraph(f"{b[1]}.  {b[2]}", style="Heading 2")
        elif k == "h3":
            doc.add_paragraph(f"{b[1]}  {b[2]}", style="Heading 3")
        elif k == "h4":
            doc.add_paragraph(f"{b[1]}. {b[2]}", style="Heading 4")
        elif k == "p":
            pp = doc.add_paragraph(); add_runs(pp, b[1])
        elif k == "bullets":
            add_bullets(doc, b[1], b[2] if len(b) > 2 else 0)
        elif k == "numbered":
            for i, it in enumerate(b[1], 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                p.paragraph_format.first_line_indent = Cm(-0.5)
                p.paragraph_format.space_after = Pt(2)
                mr = p.add_run(f"{i}.  "); mr.bold = True; mr.font.color.rgb = _rgb(T.HEADING_BLUE)
                mr.font.name = "Times New Roman"
                add_runs(p, it)
        elif k == "figure":
            add_image(doc, b[1])
            add_caption(doc, "Figure", b[2], b[3])
        elif k == "table":
            add_table(doc, b[1], b[2], b[3], b[4], b[5] if len(b) > 5 else None)
        elif k == "code":
            add_code(doc, b[1], b[2])
        elif k == "note":
            add_note(doc, b[1])
        elif k == "spacer":
            doc.add_paragraph()
        elif k == "pagebreak":
            doc.add_page_break()
        prev_break = k in ("cover", "frontmatter", "toc", "lof", "lot", "pagebreak")

    enable_update_fields(doc)
    doc.save(out_path)
    return out_path
